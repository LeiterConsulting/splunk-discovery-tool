import importlib.util
from io import BytesIO
import json
import os
import sys
import tempfile
import unittest
from pathlib import Path
from urllib.parse import parse_qs, urlsplit
import zipfile


ROOT = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))


from capabilities.install_manager import CapabilityManager
from capabilities.models import CapabilityConfig
from capabilities.rag.indexer import ArtifactSourceIndexer
from capabilities.registry import CapabilityRegistry
from capabilities.tools import DeterministicExportProvider, SplunkDeepLinkProvider, VisualizationPreviewProvider
from config_manager import ConfigManager


def build_simple_pdf_bytes(text: str) -> bytes:
    safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    objects = []
    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    objects.append(b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>\nendobj\n")
    stream = f"BT\n/F1 12 Tf\n72 720 Td\n({safe_text}) Tj\nET\n".encode("latin-1")
    objects.append(b"4 0 obj\n<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"endstream\nendobj\n")
    objects.append(b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(offsets)}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
    return bytes(pdf)


def build_simple_docx_bytes(title: str, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    from docx import Document

    buffer = BytesIO()
    document = Document()
    if title:
        document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        column_count = max(len(row) for row in table_rows if row)
        if column_count > 0:
            table = document.add_table(rows=0, cols=column_count)
            for row_values in table_rows:
                row_cells = table.add_row().cells
                for index in range(column_count):
                    row_cells[index].text = row_values[index] if index < len(row_values) else ""
    document.save(buffer)
    return buffer.getvalue()


class CapabilityFrameworkTests(unittest.TestCase):
    def test_bootstrap_persists_known_capabilities(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())

                safe_config = manager.config_manager.export_safe()
                self.assertIn("rag_local", safe_config["capabilities"])
                self.assertIn("rag_chromadb", safe_config["capabilities"])
                self.assertFalse(safe_config["capabilities"]["rag_local"]["installed"])
                self.assertEqual(safe_config["capabilities"]["rag_local"]["install_method"], "internal")

                reloaded = ConfigManager(str(config_path))
                self.assertIn("rag_local", reloaded.export_safe()["capabilities"])
            finally:
                os.chdir(original_cwd)

    def test_rag_local_install_enable_and_query_flow(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            (output_dir / "v2_insights_brief_test.md").write_text(
                "Platform health needs attention. _internal shows ingestion delays and queue pressure. "
                "Recommended next step: validate platform health and ingestion.",
                encoding="utf-8",
            )
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                config_manager = ConfigManager(str(config_path))
                manager = CapabilityManager(config_manager, registry=CapabilityRegistry())

                install_result = manager.install_capability("rag_local")
                self.assertTrue(install_result.ok)

                config_result = manager.update_capability_config(
                    "rag_local",
                    {"source_dir": str(output_dir)},
                )
                self.assertTrue(config_result.ok)

                enable_result = manager.enable_capability("rag_local")
                self.assertTrue(enable_result.ok)
                self.assertEqual(enable_result.details["status"], "ready")

                rag_result = manager.get_rag_context("What should I improve for platform health next?", max_chunks=2)

                self.assertEqual(rag_result["capability"], "rag_local")
                self.assertIn("OPTIONAL LOCAL RAG CONTEXT", rag_result["context_text"])
                self.assertGreaterEqual(len(rag_result["chunks"]), 1)

                reloaded_manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                state = reloaded_manager.get_capability_state("rag_local")

                self.assertTrue(state["installed"])
                self.assertTrue(state["enabled"])
                self.assertEqual(state["config"]["source_dir"], str(output_dir))
                self.assertEqual(state["health_status"], "ready")
                self.assertTrue(state["purpose"])
                self.assertTrue(state["intent"])
                self.assertGreaterEqual(len(state["capability_set"]), 1)
            finally:
                os.chdir(original_cwd)

    def test_enable_returns_success_when_capability_health_is_degraded(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            empty_output_dir = temp_path / "output"
            empty_output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())

                install_result = manager.install_capability("rag_local")
                self.assertTrue(install_result.ok)

                manager.update_capability_config("rag_local", {"source_dir": str(empty_output_dir)})
                enable_result = manager.enable_capability("rag_local")

                self.assertTrue(enable_result.ok)
                self.assertTrue(enable_result.state["enabled"])
                self.assertEqual(enable_result.state["health_status"], "degraded")
                self.assertIn("Capability enabled", enable_result.message or "Capability enabled")
            finally:
                os.chdir(original_cwd)

    def test_chroma_indexer_collects_typed_documents(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            (output_dir / "v2_intelligence_blueprint_test.json").write_text(
                json.dumps(
                    {
                        "readiness_score": 76,
                        "overview": {
                            "total_indexes": 47,
                            "total_sourcetypes": 24,
                            "total_hosts": 4,
                            "total_sources": 12,
                            "data_volume_24h": "~26.2GB",
                            "splunk_version": "10.0.1",
                        },
                        "recommendations": [
                            {
                                "title": "Platform Health and Splunk Operational Monitoring",
                                "priority": "high",
                                "description": "Validate ingestion, queue pressure, and scheduler health.",
                            }
                        ],
                        "coverage_gaps": [
                            {
                                "gap": "Network connectivity monitoring",
                                "priority": "medium",
                                "why_it_matters": "Packet loss and latency are not yet covered.",
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )
            (output_dir / "v2_operator_runbook_test.md").write_text(
                "## Queue Pressure\n\nOperators should validate queue pressure and ingestion delays in _internal before escalating.",
                encoding="utf-8",
            )
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                definition = manager.registry.get_definition("rag_chromadb")
                config = manager.config_manager.get_capability("rag_chromadb")
                indexer = ArtifactSourceIndexer(config=config, definition=definition)
                documents = indexer.collect_documents()

                self.assertGreaterEqual(len(documents), 3)
                source_types = {document.source_type for document in documents}
                self.assertIn("discovery_artifact", source_types)
                self.assertIn("runbook", source_types)
                self.assertTrue(any("Platform Health" in document.content for document in documents))
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_reindex_and_query_flow(self):
        if importlib.util.find_spec("chromadb") is None:
            self.skipTest("chromadb is not installed in the active environment")

        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            (output_dir / "v2_operator_runbook_test.md").write_text(
                "Queue pressure and ingestion delays require validation in _internal. Scheduler lag and blocked pipelines should be checked next.",
                encoding="utf-8",
            )
            (output_dir / "v2_ai_summary_test.json").write_text(
                json.dumps(
                    {
                        "ai_summary": "The latest artifacts call out queue pressure, ingestion delay, and scheduler backlog as the top operational concerns."
                    }
                ),
                encoding="utf-8",
            )
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())

                install_result = manager.install_capability("rag_chromadb")
                self.assertTrue(install_result.ok)

                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                enable_result = manager.enable_capability("rag_chromadb")
                self.assertTrue(enable_result.ok)
                self.assertTrue(enable_result.state["enabled"])

                reindex_result = manager.reindex_capability("rag_chromadb")
                self.assertTrue(reindex_result.ok)
                self.assertGreater(reindex_result.details["index_summary"]["document_count"], 0)

                test_result = manager.test_capability("rag_chromadb")
                self.assertTrue(test_result.ok)
                self.assertEqual(test_result.state["health_status"], "ready")

                rag_result = manager.get_rag_context("What do the artifacts say about queue pressure and ingestion delays?", max_chunks=2)

                self.assertEqual(rag_result["capability"], "rag_chromadb")
                self.assertIn("OPTIONAL CHROMADB RAG CONTEXT", rag_result["context_text"])
                self.assertGreaterEqual(len(rag_result["chunks"]), 1)
                self.assertIn("source_type", rag_result["chunks"][0]["metadata"])
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_managed_asset_import_and_listing_flow(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                import_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Indexer cluster dependencies",
                        "asset_type": "connected_system_context",
                        "source_label": "Production Splunk Platform",
                        "description": "Connectivity and service dependencies for the production indexer tier.",
                        "tags": ["splunk", "network", "dependencies"],
                        "content": "Indexers receive data from heavy forwarders, rely on shared storage, and forward alerts to external paging systems.",
                    },
                )

                self.assertTrue(import_result.ok)
                self.assertEqual(import_result.details["asset"]["asset_type"], "connected_system_context")
                self.assertFalse(import_result.details["auto_reindexed"])
                self.assertGreaterEqual(len(import_result.details["asset"]["focus_terms"]), 1)
                self.assertGreaterEqual(len(import_result.details["asset"]["key_points"]), 1)
                self.assertGreaterEqual(len(import_result.details["asset"]["usage_guidance"]), 1)

                list_result = manager.list_rag_assets("rag_chromadb")
                self.assertTrue(list_result.ok)
                self.assertEqual(list_result.details["asset_count"], 1)
                self.assertEqual(list_result.details["assets"][0]["title"], "Indexer cluster dependencies")
                self.assertGreaterEqual(len(list_result.details["assets"][0]["focus_terms"]), 1)

                capability_state = manager.get_capability_state("rag_chromadb")
                self.assertEqual(capability_state["knowledge_asset_summary"]["asset_count"], 1)

                reloaded_manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                reloaded_assets = reloaded_manager.list_rag_assets("rag_chromadb")
                self.assertTrue(reloaded_assets.ok)
                self.assertEqual(reloaded_assets.details["asset_count"], 1)
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_spl_library_asset_persists_structured_attributes(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )
                (output_dir / "discovery_sessions.json").write_text(
                    json.dumps([
                        {
                            "timestamp": "20260514_120000",
                            "created_at": "2026-05-14T12:00:00",
                            "overview": {"splunk_version": "10.0.1"},
                            "mcp_capabilities": {"tools": ["splunk_run_query"]},
                            "report_paths": ["v2_intelligence_blueprint_20260514_120000.json"],
                            "readiness_score": 92,
                        }
                    ]),
                    encoding="utf-8",
                )
                (output_dir / "v2_intelligence_blueprint_20260514_120000.json").write_text(
                    json.dumps(
                        {
                            "generated_at": "2026-05-14T12:00:00",
                            "readiness_score": 92,
                            "overview": {"splunk_version": "10.0.1"},
                            "finding_ledger": [
                                {"title": "Analyzing index: _internal", "data": {"title": "_internal", "totalEventCount": "12"}},
                                {"title": "Analyzing sourcetype: splunkd", "data": {"sourcetype": "splunkd"}},
                            ],
                        }
                    ),
                    encoding="utf-8",
                )

                import_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Saved SPL Library Query",
                        "asset_type": "spl_query_library",
                        "source_label": "Chat assistant",
                        "description": "Saved reusable SPL for later chat and Splunk execution.",
                        "tags": ["spl", "library", "chat"],
                        "content": "Saved SPL query for reuse. Query summary: search index=_internal | stats count by sourcetype.",
                        "attributes": {
                            "spl_query": "search index=_internal | stats count by sourcetype",
                            "app": "search",
                            "earliest": "-24h",
                            "latest": "now",
                            "origin_kind": "chat_assistant",
                        },
                    },
                )

                self.assertTrue(import_result.ok)
                self.assertEqual(import_result.details["asset"]["asset_type"], "spl_query_library")
                self.assertEqual(
                    import_result.details["asset"]["attributes"]["spl_query"],
                    "search index=_internal | stats count by sourcetype",
                )
                self.assertEqual(import_result.details["asset"]["attributes"]["origin_kind"], "chat_assistant")
                self.assertEqual(
                    import_result.details["asset"]["attributes"]["spl_intelligence"]["indexes"],
                    ["_internal"],
                )
                self.assertEqual(
                    import_result.details["asset"]["attributes"]["spl_intelligence"]["query_intent"],
                    "inventory_aggregation",
                )
                self.assertEqual(
                    import_result.details["asset"]["attributes"]["spl_intelligence"]["environment_fit"]["status"],
                    "strong",
                )
                self.assertTrue(
                    import_result.details["asset"]["attributes"]["spl_intelligence"]["environment_fit"]["score"] >= 80
                )
                self.assertEqual(
                    import_result.details["asset"]["attributes"]["spl_intelligence"]["validation"]["status"],
                    "unvalidated",
                )
                self.assertEqual(
                    import_result.details["asset"]["attributes"]["spl_intelligence"]["reuse"]["tier"],
                    "preferred",
                )

                list_result = manager.list_rag_assets("rag_chromadb")
                self.assertTrue(list_result.ok)
                self.assertEqual(list_result.details["asset_count"], 1)
                self.assertEqual(list_result.details["asset_type_counts"]["spl_query_library"], 1)
                self.assertEqual(
                    list_result.details["assets"][0]["attributes"]["spl_query"],
                    "search index=_internal | stats count by sourcetype",
                )

                detail_result = manager.get_rag_asset_detail(
                    "rag_chromadb",
                    import_result.details["asset"]["asset_id"],
                )
                self.assertTrue(detail_result.ok)
                self.assertEqual(
                    detail_result.details["asset"]["attributes"]["app"],
                    "search",
                )
                self.assertEqual(
                    detail_result.details["asset"]["attributes"]["spl_intelligence"]["environment_fit"]["splunk_version"],
                    "10.0.1",
                )

                reloaded_manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                reloaded_assets = reloaded_manager.list_rag_assets("rag_chromadb")
                self.assertTrue(reloaded_assets.ok)
                self.assertEqual(
                    reloaded_assets.details["assets"][0]["attributes"]["spl_query"],
                    "search index=_internal | stats count by sourcetype",
                )
                self.assertEqual(
                    reloaded_assets.details["assets"][0]["attributes"]["spl_intelligence"]["environment_fit"]["status"],
                    "strong",
                )
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_managed_asset_library_check_out_and_check_in_flow(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                import_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "SOAR integration routing notes",
                        "asset_type": "integration_context",
                        "source_label": "SOAR platform",
                        "description": "Routing dependencies for the SOAR handoff path.",
                        "tags": ["soar", "routing", "handoff"],
                        "content": "SOAR alert escalation depends on the webhook relay, the shared certificate rotation path, and the on-call routing roster.",
                    },
                )

                self.assertTrue(import_result.ok)
                asset_id = import_result.details["asset"]["asset_id"]
                self.assertEqual(import_result.details["asset"]["library_status"], "checked_in")
                self.assertTrue(import_result.details["asset"].get("last_checked_in_at"))
                self.assertFalse(import_result.details["asset"].get("checked_out_at"))

                list_result = manager.list_rag_assets("rag_chromadb")
                self.assertTrue(list_result.ok)
                self.assertEqual(list_result.details["asset_count"], 1)
                self.assertEqual(list_result.details["checked_in_asset_count"], 1)
                self.assertEqual(list_result.details["checked_out_asset_count"], 0)
                self.assertEqual(list_result.details["library_status_counts"]["checked_in"], 1)

                definition = manager.registry.get_definition("rag_chromadb")
                config = manager.config_manager.get_capability("rag_chromadb")
                indexer = ArtifactSourceIndexer(config=config, definition=definition)
                checked_in_documents = [
                    document for document in indexer.collect_documents() if document.source_type == "knowledge_asset"
                ]
                self.assertGreaterEqual(len(checked_in_documents), 1)
                self.assertTrue(all(document.metadata.get("asset_library_status") == "checked_in" for document in checked_in_documents))

                check_out_result = manager.check_out_rag_asset("rag_chromadb", asset_id)
                self.assertTrue(check_out_result.ok)
                self.assertTrue(check_out_result.details["changed"])
                self.assertEqual(check_out_result.details["asset"]["library_status"], "checked_out")
                self.assertTrue(check_out_result.details["asset"].get("checked_out_at"))
                self.assertEqual(check_out_result.details["asset_summary"]["checked_in_asset_count"], 0)
                self.assertEqual(check_out_result.details["asset_summary"]["checked_out_asset_count"], 1)

                checked_out_documents = [
                    document for document in ArtifactSourceIndexer(config=config, definition=definition).collect_documents()
                    if document.source_type == "knowledge_asset"
                ]
                self.assertEqual(len(checked_out_documents), 0)

                check_in_result = manager.check_in_rag_asset("rag_chromadb", asset_id)
                self.assertTrue(check_in_result.ok)
                self.assertTrue(check_in_result.details["changed"])
                self.assertEqual(check_in_result.details["asset"]["library_status"], "checked_in")
                self.assertTrue(check_in_result.details["asset"].get("last_checked_in_at"))
                self.assertEqual(check_in_result.details["asset_summary"]["checked_in_asset_count"], 1)
                self.assertEqual(check_in_result.details["asset_summary"]["checked_out_asset_count"], 0)

                checked_back_in_documents = [
                    document for document in ArtifactSourceIndexer(config=config, definition=definition).collect_documents()
                    if document.source_type == "knowledge_asset"
                ]
                self.assertGreaterEqual(len(checked_back_in_documents), 1)
                self.assertTrue(all(document.metadata.get("asset_library_status") == "checked_in" for document in checked_back_in_documents))

                capability_state = manager.get_capability_state("rag_chromadb")
                self.assertEqual(capability_state["knowledge_asset_summary"]["checked_in_asset_count"], 1)
                self.assertEqual(capability_state["knowledge_asset_summary"]["checked_out_asset_count"], 0)
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_spl_library_duplicate_save_refreshes_existing_asset(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                initial_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Saved SPL Library Query",
                        "asset_type": "spl_query_library",
                        "source_label": "Chat assistant",
                        "description": "Original reusable SPL entry.",
                        "tags": ["spl", "library", "chat"],
                        "content": "Saved SPL query for reuse.\n\n## Query\nsearch index=_internal | stats count by sourcetype",
                        "attributes": {
                            "spl_query": "search index=_internal | stats count by sourcetype",
                            "app": "search",
                            "earliest": "-24h",
                            "latest": "now",
                            "origin_kind": "chat_assistant",
                        },
                    },
                )

                self.assertTrue(initial_result.ok)
                self.assertEqual(initial_result.details["asset_import_action"], "created")
                initial_asset = initial_result.details["asset"]

                refreshed_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Top Sourcetypes SPL",
                        "asset_type": "spl_query_library",
                        "source_label": "Report viewer",
                        "description": "Refreshed metadata for the same reusable SPL entry.",
                        "tags": ["spl", "library", "report"],
                        "content": "Saved SPL query for reuse.\n\n## Query\nsearch   index=_internal   |   stats count by sourcetype\n\n## Context\nDetected in report output.",
                        "attributes": {
                            "spl_query": "search   index=_internal   |   stats count by sourcetype",
                            "app": "search",
                            "earliest": "-7d",
                            "latest": "now",
                            "origin_kind": "report_viewer",
                            "origin_label": "summary_report.md",
                        },
                    },
                )

                self.assertTrue(refreshed_result.ok)
                self.assertEqual(refreshed_result.message, "Knowledge asset refreshed. Enable indexed retrieval to use it in context previews and chat.")
                self.assertEqual(refreshed_result.details["asset_import_action"], "updated")
                refreshed_asset = refreshed_result.details["asset"]
                self.assertEqual(refreshed_asset["asset_id"], initial_asset["asset_id"])
                self.assertEqual(refreshed_asset["content_path"], initial_asset["content_path"])
                self.assertEqual(refreshed_asset["created_at"], initial_asset["created_at"])
                self.assertGreaterEqual(refreshed_asset["updated_at"], initial_asset["updated_at"])
                self.assertEqual(refreshed_asset["title"], "Top Sourcetypes SPL")
                self.assertEqual(refreshed_asset["source_label"], "Report viewer")
                self.assertEqual(refreshed_asset["attributes"]["earliest"], "-7d")
                self.assertEqual(refreshed_asset["attributes"]["origin_kind"], "report_viewer")

                list_result = manager.list_rag_assets("rag_chromadb")
                self.assertTrue(list_result.ok)
                self.assertEqual(list_result.details["asset_count"], 1)
                self.assertEqual(list_result.details["asset_type_counts"]["spl_query_library"], 1)
                self.assertEqual(list_result.details["assets"][0]["title"], "Top Sourcetypes SPL")

                detail_result = manager.get_rag_asset_detail("rag_chromadb", refreshed_asset["asset_id"])
                self.assertTrue(detail_result.ok)
                self.assertIn("Detected in report output.", detail_result.details["context_body"])
                self.assertEqual(detail_result.details["asset"]["attributes"]["origin_label"], "summary_report.md")

                reloaded_manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                reloaded_assets = reloaded_manager.list_rag_assets("rag_chromadb")
                self.assertTrue(reloaded_assets.ok)
                self.assertEqual(reloaded_assets.details["asset_count"], 1)
                self.assertEqual(reloaded_assets.details["assets"][0]["asset_id"], initial_asset["asset_id"])
                self.assertEqual(reloaded_assets.details["assets"][0]["attributes"]["earliest"], "-7d")
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_asset_detail_exposes_stored_sections_and_chunk_browser(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                import_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Forwarder trust guidance",
                        "asset_type": "connected_system_context",
                        "source_label": "Forwarder trust path",
                        "description": "Support note for certificate and dependency review.",
                        "tags": ["forwarders", "certificates"],
                        "content": "## Dependency Summary\n- Universal forwarders depend on certificate trust and the indexer path.\n- Operators should review certificate status before escalating.\n",
                    },
                )
                self.assertTrue(import_result.ok)

                detail_result = manager.get_rag_asset_detail(
                    "rag_chromadb",
                    import_result.details["asset"]["asset_id"],
                )

                self.assertTrue(detail_result.ok)
                self.assertEqual(detail_result.details["asset"]["title"], "Forwarder trust guidance")
                self.assertGreaterEqual(len(detail_result.details["stored_sections"]), 3)
                self.assertTrue(any(section["title"] == "Dependency Summary" for section in detail_result.details["stored_sections"]))
                self.assertIn("certificate trust", detail_result.details["context_body"])
                self.assertGreaterEqual(len(detail_result.details["chunk_sections"]), 1)
                self.assertEqual(
                    detail_result.details["chunk_sections"][0]["metadata"]["asset_type"],
                    "connected_system_context",
                )
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_spl_library_feedback_updates_validation_state(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )
                (output_dir / "discovery_sessions.json").write_text(
                    json.dumps([
                        {
                            "timestamp": "20260514_120000",
                            "created_at": "2026-05-14T12:00:00",
                            "overview": {"splunk_version": "10.0.1"},
                            "mcp_capabilities": {"tools": ["splunk_run_query"]},
                            "report_paths": ["v2_intelligence_blueprint_20260514_120000.json"],
                            "readiness_score": 92,
                        }
                    ]),
                    encoding="utf-8",
                )
                (output_dir / "v2_intelligence_blueprint_20260514_120000.json").write_text(
                    json.dumps(
                        {
                            "generated_at": "2026-05-14T12:00:00",
                            "readiness_score": 92,
                            "overview": {"splunk_version": "10.0.1"},
                            "finding_ledger": [
                                {"title": "Analyzing index: _internal", "data": {"title": "_internal", "totalEventCount": "12"}},
                            ],
                        }
                    ),
                    encoding="utf-8",
                )

                import_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Saved SPL Library Query",
                        "asset_type": "spl_query_library",
                        "source_label": "Chat assistant",
                        "description": "Saved reusable SPL for later chat and Splunk execution.",
                        "tags": ["spl", "library", "chat"],
                        "content": "Saved SPL query for reuse.",
                        "attributes": {
                            "spl_query": "search index=_internal | stats count by sourcetype",
                            "app": "search",
                            "earliest": "-24h",
                            "latest": "now",
                            "origin_kind": "chat_assistant",
                        },
                    },
                )

                self.assertTrue(import_result.ok)

                success_feedback = manager.record_rag_spl_query_feedback(
                    "rag_chromadb",
                    "search index=_internal | stats count by sourcetype",
                    "success",
                    {
                        "row_count": 4,
                        "earliest_time": "-24h",
                        "latest_time": "now",
                    },
                )
                self.assertTrue(success_feedback.ok)
                self.assertEqual(
                    success_feedback.details["asset"]["attributes"]["spl_intelligence"]["validation"]["status"],
                    "known_good",
                )
                self.assertEqual(
                    success_feedback.details["asset"]["attributes"]["spl_intelligence"]["validation"]["success_count"],
                    1,
                )
                self.assertEqual(
                    success_feedback.details["asset"]["attributes"]["spl_intelligence"]["reuse"]["tier"],
                    "known_good",
                )

                failure_feedback = manager.record_rag_spl_query_feedback(
                    "rag_chromadb",
                    "search index=_internal | stats count by sourcetype",
                    "failure",
                    {
                        "error": "Search failed due to permissions.",
                        "earliest_time": "-24h",
                        "latest_time": "now",
                    },
                )
                self.assertTrue(failure_feedback.ok)
                self.assertEqual(
                    failure_feedback.details["asset"]["attributes"]["spl_intelligence"]["validation"]["status"],
                    "mixed",
                )
                self.assertEqual(
                    failure_feedback.details["asset"]["attributes"]["spl_intelligence"]["validation"]["failure_count"],
                    1,
                )
                self.assertEqual(
                    failure_feedback.details["asset"]["attributes"]["spl_intelligence"]["validation"]["last_error"],
                    "Search failed due to permissions.",
                )

                reloaded_manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                reloaded_assets = reloaded_manager.list_rag_assets("rag_chromadb")
                self.assertTrue(reloaded_assets.ok)
                self.assertEqual(
                    reloaded_assets.details["assets"][0]["attributes"]["spl_intelligence"]["validation"]["execution_count"],
                    2,
                )
                self.assertEqual(
                    reloaded_assets.details["assets"][0]["attributes"]["spl_intelligence"]["validation"]["status"],
                    "mixed",
                )
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_search_surfaces_known_good_reusable_spl_queries(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )
                (output_dir / "discovery_sessions.json").write_text(
                    json.dumps([
                        {
                            "timestamp": "20260514_120000",
                            "created_at": "2026-05-14T12:00:00",
                            "overview": {"splunk_version": "10.0.1"},
                            "mcp_capabilities": {"tools": ["splunk_run_query"]},
                            "report_paths": ["v2_intelligence_blueprint_20260514_120000.json"],
                            "readiness_score": 92,
                        }
                    ]),
                    encoding="utf-8",
                )
                (output_dir / "v2_intelligence_blueprint_20260514_120000.json").write_text(
                    json.dumps(
                        {
                            "generated_at": "2026-05-14T12:00:00",
                            "readiness_score": 92,
                            "overview": {"splunk_version": "10.0.1"},
                            "finding_ledger": [
                                {"title": "Analyzing index: _internal", "data": {"title": "_internal", "totalEventCount": "12"}},
                                {"title": "Analyzing index: netops", "data": {"title": "netops", "totalEventCount": "12"}},
                            ],
                        }
                    ),
                    encoding="utf-8",
                )

                manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Internal Sourcetype Counts",
                        "asset_type": "spl_query_library",
                        "source_label": "Chat assistant",
                        "description": "Check top sourcetypes in internal logs.",
                        "tags": ["spl", "internal"],
                        "content": "Reusable SPL for internal telemetry analysis.",
                        "attributes": {
                            "spl_query": "search index=_internal | stats count by sourcetype",
                            "app": "search",
                            "earliest": "-24h",
                            "latest": "now",
                        },
                    },
                )
                manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Netops Host Counts",
                        "asset_type": "spl_query_library",
                        "source_label": "Chat assistant",
                        "description": "Check top hosts in network operations data.",
                        "tags": ["spl", "netops"],
                        "content": "Reusable SPL for network operations telemetry.",
                        "attributes": {
                            "spl_query": "search index=netops | stats count by host",
                            "app": "search",
                            "earliest": "-24h",
                            "latest": "now",
                        },
                    },
                )
                manager.record_rag_spl_query_feedback(
                    "rag_chromadb",
                    "search index=_internal | stats count by sourcetype",
                    "success",
                    {"row_count": 3, "earliest_time": "-24h", "latest_time": "now"},
                )

                definition = manager.registry.get_definition("rag_chromadb")
                config = manager.config_manager.get_capability("rag_chromadb")
                indexer = ArtifactSourceIndexer(config=config, definition=definition)
                indexer.reindex()
                search_result = indexer.search("Help me reuse a query to inspect internal sourcetype counts.", max_chunks=3)

                self.assertTrue(search_result["reusable_spl_queries"])
                self.assertEqual(
                    search_result["reusable_spl_queries"][0]["query"],
                    "search index=_internal | stats count by sourcetype",
                )
                self.assertTrue(search_result["reusable_spl_queries"][0]["known_good"])
                self.assertIn("REUSABLE SPL QUERY CANDIDATES", search_result["context_text"])
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_search_migrates_legacy_saved_spl_library_assets(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                import_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "SPL Library: | tstats count where index=* by index | sort - count | head 25",
                        "asset_type": "reference_document",
                        "source_label": "Chat assistant response",
                        "description": "Saved reusable SPL query for direct Splunk launch and chat reuse.",
                        "tags": ["spl", "spl-library", "chat-assistant"],
                        "content": (
                            "Saved SPL query for reuse in Splunk Web and follow-on chat workflows.\n"
                            "Query summary: | tstats count where index=* by index | sort - count | head 25\n"
                            "Saved from: Chat assistant response."
                        ),
                        "attributes": {},
                    },
                )
                self.assertTrue(import_result.ok)
                self.assertEqual(import_result.details["asset"]["asset_type"], "reference_document")

                definition = manager.registry.get_definition("rag_chromadb")
                config = manager.config_manager.get_capability("rag_chromadb")
                indexer = ArtifactSourceIndexer(config=config, definition=definition)
                indexer.reindex()

                migrated_assets = manager.list_rag_assets("rag_chromadb")
                self.assertTrue(migrated_assets.ok)
                self.assertEqual(migrated_assets.details["asset_type_counts"]["spl_query_library"], 1)
                self.assertEqual(
                    migrated_assets.details["assets"][0]["attributes"]["spl_query"],
                    "| tstats count where index=* by index | sort - count | head 25",
                )

                search_result = indexer.search("top indexes by event count", max_chunks=3)

                self.assertTrue(search_result["reusable_spl_queries"])
                self.assertEqual(
                    search_result["reusable_spl_queries"][0]["query"],
                    "| tstats count where index=* by index | sort - count | head 25",
                )
                self.assertIn("REUSABLE SPL QUERY CANDIDATES", search_result["context_text"])
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_import_file_asset_supports_pdf_uploads(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                import_result = manager.import_rag_file_asset(
                    "rag_chromadb",
                    filename="forwarder-trust.pdf",
                    content_bytes=build_simple_pdf_bytes("Forwarders need certificate trust before data reaches the indexer tier."),
                    payload={
                        "title": "Forwarder PDF guidance",
                        "asset_type": "connected_system_context",
                        "source_label": "Forwarder PDF",
                        "description": "PDF validation asset",
                        "tags": ["forwarders", "pdf"],
                    },
                )

                self.assertTrue(import_result.ok)
                self.assertEqual(import_result.details["asset"]["original_filename"], "forwarder-trust.pdf")
                self.assertEqual(import_result.details["asset"]["import_method"], "file_upload")
                self.assertIn("Forwarders need certificate trust", import_result.details["asset"]["summary"])

                detail_result = manager.get_rag_asset_detail(
                    "rag_chromadb",
                    import_result.details["asset"]["asset_id"],
                )
                self.assertTrue(detail_result.ok)
                self.assertIn("PDF Page 1", detail_result.details["context_body"])
                self.assertTrue(any(section["title"] == "PDF Page 1" for section in detail_result.details["stored_sections"]))
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_import_file_asset_supports_docx_uploads(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                import_result = manager.import_rag_file_asset(
                    "rag_chromadb",
                    filename="forwarder-escalation.docx",
                    content_bytes=build_simple_docx_bytes(
                        "Forwarder Escalation",
                        [
                            "Verify certificate trust before escalating indexer path issues.",
                            "Confirm the forwarding queue is healthy before involving platform operations.",
                        ],
                        table_rows=[
                            ["Owner", "Platform Operations"],
                            ["Queue", "Forwarding Pipeline"],
                        ],
                    ),
                    payload={
                        "title": "Forwarder DOCX guidance",
                        "asset_type": "runbook_context",
                        "source_label": "Forwarder DOCX",
                        "description": "DOCX validation asset",
                        "tags": ["forwarders", "docx"],
                    },
                )

                self.assertTrue(import_result.ok)
                self.assertEqual(import_result.details["asset"]["original_filename"], "forwarder-escalation.docx")
                self.assertEqual(import_result.details["asset"]["import_method"], "file_upload")
                self.assertIn("Verify certificate trust", import_result.details["asset"]["summary"])

                detail_result = manager.get_rag_asset_detail(
                    "rag_chromadb",
                    import_result.details["asset"]["asset_id"],
                )
                self.assertTrue(detail_result.ok)
                self.assertIn("Forwarder Escalation", detail_result.details["context_body"])
                self.assertTrue(any(section["title"] == "Forwarder Escalation" for section in detail_result.details["stored_sections"]))
                self.assertTrue(any(section["title"] == "DOCX Table 1" for section in detail_result.details["stored_sections"]))
            finally:
                os.chdir(original_cwd)

    def test_chroma_indexer_collects_knowledge_asset_metadata(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )
                manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Updated Splunk auth notes",
                        "asset_type": "splunk_documentation",
                        "source_label": "Internal documentation",
                        "tags": ["splunk", "authentication"],
                        "content": "Search head authentication now depends on the enterprise IdP group sync job and the secret rotation schedule.",
                    },
                )

                definition = manager.registry.get_definition("rag_chromadb")
                config = manager.config_manager.get_capability("rag_chromadb")
                indexer = ArtifactSourceIndexer(config=config, definition=definition)
                documents = indexer.collect_documents()

                knowledge_documents = [document for document in documents if document.source_type == "knowledge_asset"]
                self.assertGreaterEqual(len(knowledge_documents), 1)
                self.assertTrue(any(document.metadata.get("asset_id") for document in knowledge_documents))
                self.assertTrue(any(document.metadata.get("asset_type") == "splunk_documentation" for document in knowledge_documents))
                self.assertTrue(any(document.metadata.get("asset_focus_terms") for document in knowledge_documents))
            finally:
                os.chdir(original_cwd)

    def test_rag_chromadb_context_preview_and_delete_flow_for_managed_assets(self):
        if importlib.util.find_spec("chromadb") is None:
            self.skipTest("chromadb is not installed in the active environment")

        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())

                install_result = manager.install_capability("rag_chromadb")
                self.assertTrue(install_result.ok)

                manager.update_capability_config(
                    "rag_chromadb",
                    {
                        "source_dir": str(output_dir),
                        "storage_dir": str(output_dir / "rag" / "chromadb"),
                    },
                )

                enable_result = manager.enable_capability("rag_chromadb")
                self.assertTrue(enable_result.ok)

                import_result = manager.import_rag_text_asset(
                    "rag_chromadb",
                    {
                        "title": "Payments platform integration notes",
                        "asset_type": "integration_context",
                        "source_label": "Payments API",
                        "tags": ["payments", "api", "splunk"],
                        "content": "The payments API sends structured JSON events through heavy forwarders, depends on the shared certificate rotation process, and escalates through the payments support group roster.",
                    },
                )
                self.assertTrue(import_result.ok)
                self.assertTrue(import_result.details["auto_reindexed"])

                asset_path = output_dir / "rag" / "assets" / import_result.details["asset"]["content_path"]
                stored_asset = asset_path.read_text(encoding="utf-8")
                self.assertIn("## Focus Terms", stored_asset)
                self.assertIn("## Key Points", stored_asset)
                before_focus_terms, after_focus_terms = stored_asset.split("## Focus Terms", 1)
                _, after_key_points_marker = after_focus_terms.split("## Key Points", 1)
                asset_path.write_text(
                    f"{before_focus_terms}## Focus Terms\n\n- the\n- payments\n\n## Key Points{after_key_points_marker}",
                    encoding="utf-8",
                )

                preview_result = manager.build_rag_context_preview(
                    "rag_chromadb",
                    "What does Splunk need to know about the payments API integration?",
                    max_chunks=3,
                )
                self.assertTrue(preview_result.ok)
                self.assertIn("Knowledge asset context preview", preview_result.details["context_text"])
                self.assertIn("Operator context brief", preview_result.details["operator_brief"])
                self.assertGreaterEqual(len(preview_result.details["chunks"]), 1)
                self.assertEqual(preview_result.details["chunks"][0]["metadata"].get("source_type"), "knowledge_asset")
                self.assertTrue(preview_result.details["chunks"][0]["document_id"])
                self.assertGreaterEqual(len(preview_result.details["matched_assets"]), 1)
                self.assertTrue(preview_result.details["matched_assets"][0]["why_matched"])
                self.assertGreaterEqual(len(preview_result.details["matched_assets"][0]["key_points"]), 1)
                self.assertGreaterEqual(len(preview_result.details["matched_assets"][0]["matched_chunk_ids"]), 1)
                self.assertGreaterEqual(len(preview_result.details["matched_assets"][0]["matched_chunks"]), 1)
                self.assertNotIn("the", [term.lower() for term in preview_result.details["matched_assets"][0]["focus_terms"]])
                self.assertGreaterEqual(len(preview_result.details["recommended_uses"]), 1)
                self.assertGreaterEqual(len(preview_result.details["retrieved_key_points"]), 1)
                self.assertIn(
                    "Use for ownership, escalation, and support-routing questions.",
                    preview_result.details["operator_brief"],
                )

                asset_id = import_result.details["asset"]["asset_id"]
                detail_result = manager.get_rag_asset_detail("rag_chromadb", asset_id)
                self.assertTrue(detail_result.ok)
                preview_chunk_ids = set(preview_result.details["matched_assets"][0]["matched_chunk_ids"])
                detail_chunk_ids = {section["document_id"] for section in detail_result.details["chunk_sections"]}
                self.assertTrue(preview_chunk_ids.issubset(detail_chunk_ids))

                refreshed_asset = asset_path.read_text(encoding="utf-8")
                self.assertNotIn("\n- the\n", refreshed_asset)

                delete_result = manager.delete_rag_asset("rag_chromadb", asset_id)
                self.assertTrue(delete_result.ok)
                self.assertTrue(delete_result.details["deleted"])
                self.assertEqual(delete_result.details["asset_summary"]["asset_count"], 0)
            finally:
                os.chdir(original_cwd)

    def test_deeplink_provider_derives_web_base_url_from_mcp_url(self):
        definition = CapabilityRegistry().get_definition("splunk_deeplink_tools")
        config = CapabilityConfig(
            name="splunk_deeplink_tools",
            installed=True,
            enabled=True,
            config=dict(definition.default_config),
        )

        provider = SplunkDeepLinkProvider(
            config=config,
            definition=definition,
            mcp_url="https://splunk.example.local:8089/services/mcp",
        )

        self.assertEqual(provider.resolve_web_base_url(), "https://splunk.example.local:8000")
        self.assertEqual(provider.resolve_base_url_source(), "mcp.url")

    def test_deeplink_provider_prefers_override_and_encodes_search_params(self):
        definition = CapabilityRegistry().get_definition("splunk_deeplink_tools")
        config = CapabilityConfig(
            name="splunk_deeplink_tools",
            installed=True,
            enabled=True,
            config={
                **definition.default_config,
                "web_base_url": "https://splunkweb.example.local/splunk",
                "default_app": "search",
            },
        )

        provider = SplunkDeepLinkProvider(
            config=config,
            definition=definition,
            mcp_url="https://ignored.example.local:8089/services/mcp",
        )
        deeplink = provider.build_search_link(
            "index=_internal error | stats count by host",
            earliest="-2h",
            latest="now",
        )
        parsed = urlsplit(deeplink["url"])
        params = parse_qs(parsed.query)

        self.assertEqual(f"{parsed.scheme}://{parsed.netloc}{parsed.path.rsplit('/en-US/', 1)[0]}", "https://splunkweb.example.local/splunk")
        self.assertEqual(params["q"][0], "search index=_internal error | stats count by host")
        self.assertEqual(params["earliest"][0], "-2h")
        self.assertEqual(params["latest"][0], "now")
        self.assertEqual(deeplink["base_url_source"], "capability_config.web_base_url")

    def test_deeplink_install_enable_test_and_build_flow(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.config_manager.update_mcp(url="https://splunk.example.local:8089/services/mcp")

                install_result = manager.install_capability("splunk_deeplink_tools")
                self.assertTrue(install_result.ok)

                enable_result = manager.enable_capability("splunk_deeplink_tools")
                self.assertTrue(enable_result.ok)
                self.assertEqual(enable_result.state["health_status"], "ready")

                test_result = manager.test_capability("splunk_deeplink_tools")
                self.assertTrue(test_result.ok)
                self.assertIn("sample_deeplink", test_result.details)

                build_result = manager.build_deeplink(
                    "splunk_deeplink_tools",
                    "search",
                    {
                        "query": "index=_internal | stats count by sourcetype",
                        "earliest": "-7d",
                        "latest": "now",
                    },
                )
                self.assertTrue(build_result.ok)
                deeplink = build_result.details["deeplink"]
                params = parse_qs(urlsplit(deeplink["url"]).query)

                self.assertEqual(deeplink["base_url"], "https://splunk.example.local:8000")
                self.assertEqual(params["q"][0], "search index=_internal | stats count by sourcetype")
                self.assertEqual(params["earliest"][0], "-7d")
                self.assertEqual(params["latest"][0], "now")
            finally:
                os.chdir(original_cwd)

    def test_deeplink_enable_can_succeed_while_health_is_degraded(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())
                manager.config_manager.update_mcp(url="")

                install_result = manager.install_capability("splunk_deeplink_tools")
                self.assertTrue(install_result.ok)

                enable_result = manager.enable_capability("splunk_deeplink_tools")
                self.assertTrue(enable_result.ok)
                self.assertEqual(enable_result.state["health_status"], "degraded")
                self.assertIn("base url", enable_result.message.lower())
            finally:
                os.chdir(original_cwd)

    def test_visualization_provider_builds_line_and_bar_previews(self):
        definition = CapabilityRegistry().get_definition("visualization_tools")
        config = CapabilityConfig(
            name="visualization_tools",
            installed=True,
            enabled=True,
            config=dict(definition.default_config),
        )
        provider = VisualizationPreviewProvider(config=config, definition=definition)

        line_preview = provider.build_preview(
            [
                {"_time": "2026-04-19 14:00:00.000 EDT", "count": "42"},
                {"_time": "2026-04-19 15:00:00.000 EDT", "count": "57"},
                {"_time": "2026-04-19 16:00:00.000 EDT", "count": "39"},
            ],
            payload={"query_shape": "time_series"},
        )
        bar_preview = provider.build_preview(
            [
                {"sourcetype": "WinEventLog:Security", "count": "18"},
                {"sourcetype": "XmlWinEventLog:Microsoft-Windows-Sysmon/Operational", "count": "11"},
                {"sourcetype": "splunkd", "count": "7"},
            ],
            payload={"query_shape": "aggregation"},
        )

        self.assertEqual(line_preview["chart_type"], "line")
        self.assertEqual(line_preview["x_field"], "_time")
        self.assertEqual(line_preview["y_field"], "count")
        self.assertEqual(len(line_preview["points"]), 3)
        self.assertEqual(bar_preview["chart_type"], "bar")
        self.assertEqual(bar_preview["x_field"], "sourcetype")
        self.assertEqual(bar_preview["y_field"], "count")
        self.assertGreaterEqual(len(bar_preview["points"]), 2)
        self.assertEqual(
            bar_preview["points"][1]["full_label"],
            "XmlWinEventLog:Microsoft-Windows-Sysmon/Operational",
        )
        self.assertTrue(bar_preview["points"][1]["label"].endswith("..."))

    def test_visualization_install_enable_test_and_build_flow(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())

                install_result = manager.install_capability("visualization_tools")
                self.assertTrue(install_result.ok)

                enable_result = manager.enable_capability("visualization_tools")
                self.assertTrue(enable_result.ok)
                self.assertEqual(enable_result.state["health_status"], "ready")

                test_result = manager.test_capability("visualization_tools")
                self.assertTrue(test_result.ok)
                self.assertIn("supported_chart_types", test_result.details["details"])

                build_result = manager.build_visualization(
                    "visualization_tools",
                    {
                        "rows": [
                            {"_time": "2026-04-19 14:00:00.000 EDT", "count": "42"},
                            {"_time": "2026-04-19 15:00:00.000 EDT", "count": "57"},
                            {"_time": "2026-04-19 16:00:00.000 EDT", "count": "39"},
                        ],
                        "query_shape": "time_series",
                    },
                )
                self.assertTrue(build_result.ok)
                self.assertEqual(build_result.details["visualization"]["chart_type"], "line")
                self.assertEqual(build_result.state["preview_enabled"], True)
            finally:
                os.chdir(original_cwd)

    def test_export_provider_builds_bundle_and_manifest(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            timestamp = "20260419_154500"
            (output_dir / f"v2_intelligence_blueprint_{timestamp}.json").write_text(
                json.dumps({"overview": {"total_indexes": 47, "total_sourcetypes": 24, "total_hosts": 4, "data_volume_24h": "~26.2GB"}}),
                encoding="utf-8",
            )
            (output_dir / f"v2_insights_brief_{timestamp}.md").write_text(
                "Executive brief for the latest discovery session.",
                encoding="utf-8",
            )
            (output_dir / "discovery_sessions.json").write_text(
                json.dumps(
                    [
                        {
                            "timestamp": timestamp,
                            "created_at": "2026-04-19T15:45:00",
                            "overview": {
                                "total_indexes": 47,
                                "total_sourcetypes": 24,
                                "total_hosts": 4,
                                "data_volume_24h": "~26.2GB",
                            },
                            "report_paths": [
                                f"v2_intelligence_blueprint_{timestamp}.json",
                                f"v2_insights_brief_{timestamp}.md",
                            ],
                        }
                    ]
                ),
                encoding="utf-8",
            )

            definition = CapabilityRegistry().get_definition("export_tools")
            config = CapabilityConfig(
                name="export_tools",
                installed=True,
                enabled=True,
                config={
                    **definition.default_config,
                    "source_dir": str(output_dir),
                    "export_dir": str(output_dir / "exports"),
                },
            )
            provider = DeterministicExportProvider(config=config, definition=definition)

            export_result = provider.build_export(
                {
                    "timestamp": timestamp,
                    "persona": "admin",
                    "runbook_markdown": "# Admin Runbook\n\nValidate queue pressure and ingestion delays.",
                    "runbook_filename": f"runbook_admin_{timestamp}.md",
                    "title": "Platform Health Export",
                }
            )

            zip_path = Path(export_result["bundle_path"])
            self.assertTrue(zip_path.exists())
            self.assertTrue((output_dir / "exports" / export_result["manifest_name"]).exists())
            self.assertTrue((output_dir / "exports" / export_result["summary_name"]).exists())
            self.assertEqual(export_result["artifact_count"], 2)
            self.assertIn(f"v2_intelligence_blueprint_{timestamp}.json", export_result["included_files"])

            with zipfile.ZipFile(zip_path, "r") as archive:
                names = set(archive.namelist())
                self.assertIn("manifest.json", names)
                self.assertIn("README.md", names)
                self.assertIn(f"artifacts/v2_intelligence_blueprint_{timestamp}.json", names)
                self.assertIn(f"generated/runbook_admin_{timestamp}.md", names)

    def test_export_install_enable_test_and_build_flow(self):
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
            temp_path = Path(temp_dir)
            config_path = temp_path / "config.encrypted"
            output_dir = temp_path / "output"
            output_dir.mkdir(parents=True, exist_ok=True)
            timestamp = "20260419_160000"
            (output_dir / f"v2_intelligence_blueprint_{timestamp}.json").write_text(
                json.dumps({"overview": {"total_indexes": 47, "total_sourcetypes": 24, "total_hosts": 4}}),
                encoding="utf-8",
            )
            (output_dir / f"v2_operator_runbook_{timestamp}.md").write_text(
                "Operator runbook content.",
                encoding="utf-8",
            )
            (output_dir / "discovery_sessions.json").write_text(
                json.dumps(
                    [
                        {
                            "timestamp": timestamp,
                            "created_at": "2026-04-19T16:00:00",
                            "overview": {"total_indexes": 47, "total_sourcetypes": 24, "total_hosts": 4},
                            "report_paths": [
                                f"v2_intelligence_blueprint_{timestamp}.json",
                                f"v2_operator_runbook_{timestamp}.md",
                            ],
                        }
                    ]
                ),
                encoding="utf-8",
            )
            original_cwd = Path.cwd()

            try:
                os.chdir(temp_path)
                manager = CapabilityManager(ConfigManager(str(config_path)), registry=CapabilityRegistry())

                install_result = manager.install_capability("export_tools")
                self.assertTrue(install_result.ok)

                config_result = manager.update_capability_config(
                    "export_tools",
                    {
                        "source_dir": str(output_dir),
                        "export_dir": str(output_dir / "exports"),
                    },
                )
                self.assertTrue(config_result.ok)

                enable_result = manager.enable_capability("export_tools")
                self.assertTrue(enable_result.ok)
                self.assertEqual(enable_result.state["health_status"], "ready")

                test_result = manager.test_capability("export_tools")
                self.assertTrue(test_result.ok)
                self.assertEqual(test_result.details["details"]["latest_session_timestamp"], timestamp)

                build_result = manager.build_export(
                    "export_tools",
                    {
                        "timestamp": timestamp,
                        "persona": "executive",
                        "runbook_markdown": "# Executive Runbook\n\nFocus on readiness and business impact.",
                        "runbook_filename": f"runbook_executive_{timestamp}.md",
                    },
                )
                self.assertTrue(build_result.ok)
                export_payload = build_result.details["export"]
                self.assertEqual(export_payload["session_timestamp"], timestamp)
                self.assertEqual(export_payload["persona"], "executive")
                self.assertTrue(Path(export_payload["bundle_path"]).exists())
            finally:
                os.chdir(original_cwd)


if __name__ == "__main__":
    unittest.main()